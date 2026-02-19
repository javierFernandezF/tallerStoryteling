from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import pandas as pd

# ============================================================
# CARGAR DATOS PARA REFERENCIAS EXACTAS
# ============================================================
def load_data():
    data_rows = []
    with open('/Users/javierfernandez/Desktop/Taller/Libro1.csv', 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()
    header_line = lines[0].strip()
    if header_line.startswith('"') and header_line.endswith('"'):
        header = header_line[1:-1].split(';')
    for line in lines[1:]:
        line = line.strip()
        if line.startswith('"') and line.endswith('"'):
            row_data = line[1:-1].split(';')
            data_rows.append(row_data)
    df = pd.DataFrame(data_rows, columns=header)
    numeric_cols = ['Unit price', 'Quantity', 'Tax 5%', 'Total', 'Time', 'cogs', 'gross margin percentage', 'gross income', 'Rating']
    for col in numeric_cols:
        if col in df.columns:
            df[col] = df[col].str.replace(',', '.').astype(float)
    df['Hour'] = (df['Time'] * 24).round().astype(int).clip(0, 23)
    return df

df = load_data()

# ============================================================
# CREAR DOCUMENTO WORD
# ============================================================
doc = Document()

# Configurar estilos
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================
# PORTADA
# ============================================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Desafío 3: Supermarket Sales')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Análisis de Datos para Campañas de Marketing')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph('')

client = doc.add_paragraph()
client.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = client.add_run('Cliente: Cadena de Supermercados')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# 1. DESCRIPCIÓN DEL DATASET
# ============================================================
h1 = doc.add_heading('1. Descripción del DataSet', level=1)

# Tamaño y Campos
doc.add_heading('1.1 Tamaño y Campos', level=2)

p = doc.add_paragraph()
p.add_run('Tamaño del dataset: ').bold = True
p.add_run(f'{len(df)} registros (filas) y {len(df.columns)} campos (columnas).')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Campos del dataset:').bold = True

# Tabla de campos
table = doc.add_table(rows=18, cols=3, style='Light Grid Accent 1')
table.columns[0].width = Cm(4)
table.columns[1].width = Cm(3)
table.columns[2].width = Cm(9)

# Header
headers = ['Campo', 'Tipo', 'Descripción']
for i, header_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header_text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data
fields = [
    ('Invoice ID', 'Texto', 'Identificador único de la factura'),
    ('Branch', 'Texto', 'Sucursal (A, B o C)'),
    ('City', 'Texto', 'Ciudad (Yangon, Mandalay, Naypyitaw)'),
    ('Customer type', 'Texto', 'Tipo de cliente (Member o Normal)'),
    ('Gender', 'Texto', 'Género del cliente (Male o Female)'),
    ('Product line', 'Texto', 'Línea de producto (6 categorías)'),
    ('Unit price', 'Numérico', 'Precio unitario del producto'),
    ('Quantity', 'Numérico', 'Cantidad comprada'),
    ('Tax 5%', 'Numérico', 'Impuesto del 5% sobre la compra'),
    ('Total', 'Numérico', 'Monto total de la transacción'),
    ('Date', 'Fecha', 'Fecha de la transacción'),
    ('Time', 'Numérico', 'Hora de la transacción'),
    ('Payment', 'Texto', 'Método de pago (Ewallet, Cash, Credit card)'),
    ('cogs', 'Numérico', 'Costo de los bienes vendidos'),
    ('gross margin percentage', 'Numérico', 'Porcentaje de margen bruto'),
    ('gross income', 'Numérico', 'Ingreso bruto'),
    ('Rating', 'Numérico', 'Calificación del cliente (1-10)'),
]

for i, (field, tipo, desc) in enumerate(fields):
    table.rows[i+1].cells[0].text = field
    table.rows[i+1].cells[1].text = tipo
    table.rows[i+1].cells[2].text = desc

# Principales Observaciones
doc.add_heading('1.2 Principales Observaciones del DataSet', level=2)

observations = [
    f'El dataset contiene {len(df)} transacciones de ventas de una cadena de supermercados.',
    'Los datos abarcan 3 meses de operación.',
    f'Las ventas totales registradas son de ${df["Total"].sum():,.2f}.',
    f'El ticket promedio por transacción es de ${df["Total"].mean():.2f}.',
    f'Existen 3 sucursales: A (Yangon) con {len(df[df["Branch"]=="A"])} transacciones, B (Mandalay) con {len(df[df["Branch"]=="B"])} transacciones y C (Naypyitaw) con {len(df[df["Branch"]=="C"])} transacciones.',
    f'La distribución por género es equilibrada: {len(df[df["Gender"]=="Female"])} mujeres y {len(df[df["Gender"]=="Male"])} hombres.',
    f'Hay 2 tipos de clientes: Member ({len(df[df["Customer type"]=="Member"])}) y Normal ({len(df[df["Customer type"]=="Normal"])}).',
    'Existen 6 líneas de productos: Health and beauty, Electronic accessories, Home and lifestyle, Sports and travel, Food and beverages, y Fashion accessories.',
    'El horario de operación va de 10:00 a 21:00 hrs (12 horas).',
    '3 métodos de pago: Ewallet, Cash y Credit card.',
]

for obs in observations:
    doc.add_paragraph(obs, style='List Bullet')

doc.add_page_break()

# ============================================================
# 2. CONTEXTO
# ============================================================
doc.add_heading('2. Contexto', level=1)

doc.add_heading('2.1 ¿Quién?', level=2)
p = doc.add_paragraph()
p.add_run('Cliente: ').bold = True
p.add_run('Cadena de Supermercados con varias sucursales en distintas ciudades de un determinado país.')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Sucursales:').bold = True

sucursales = [
    'Sucursal A - Yangon (340 transacciones)',
    'Sucursal B - Mandalay (332 transacciones)',
    'Sucursal C - Naypyitaw (328 transacciones)',
]
for s in sucursales:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('2.2 ¿Qué?', level=2)
p = doc.add_paragraph()
p.add_run('Objetivo del negocio: ').bold = True
p.add_run('La empresa quiere realizar 3 campañas de Marketing, una por mes, con el objetivo de maximizar los ingresos por las ventas.')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Preguntas guía del cliente:').bold = True

preguntas = [
    '¿Cómo son las compras de Mujeres comparadas con las de los Hombres?',
    '¿Cómo se distribuyen las compras en las distintas horas del día?',
    '¿Cómo son los ingresos por línea de productos?',
    '¿Cómo son los Ingresos por línea de productos y sexo?',
    '¿Cómo son las compras en las distintas Sucursales?',
    '¿Cómo son las compras por línea de productos?',
]
for preg in preguntas:
    doc.add_paragraph(preg, style='List Bullet')

doc.add_page_break()

# ============================================================
# 3. STORYBOARD
# ============================================================
doc.add_heading('3. Storyboard', level=1)

p = doc.add_paragraph()
p.add_run('A continuación se describe el storyboard planificado para la presentación, slide por slide.').italic = True

# Slide 1: Introducción
doc.add_heading('3.1 Introducción (Slide 1)', level=2)
p = doc.add_paragraph()
p.add_run('Contenido de la slide:').bold = True

intro_items = [
    'Título: "Análisis de Ventas - Cadena de Supermercados"',
    'Contexto: La cadena de supermercados opera en 3 ciudades y busca maximizar ingresos a través de 3 campañas de marketing mensuales.',
    f'Datos clave: {len(df)} transacciones analizadas, ${df["Total"].sum():,.2f} en ventas totales, 3 meses de datos.',
    'Pregunta central: ¿Dónde están las oportunidades para aumentar las ventas?',
]
for item in intro_items:
    doc.add_paragraph(item, style='List Bullet')

# Slides de Objetivo
doc.add_heading('3.2 Objetivo de la presentación (Slides 2-7)', level=2)

# Slide 2
p = doc.add_paragraph()
run = p.add_run('Slide 2: Compras por Género')
run.bold = True

female_sales = df[df['Gender'] == 'Female']['Total'].sum()
male_sales = df[df['Gender'] == 'Male']['Total'].sum()
total_sales = female_sales + male_sales

slide2_items = [
    f'Mujeres: ${female_sales:,.2f} ({female_sales/total_sales*100:.1f}%) - {len(df[df["Gender"]=="Female"])} transacciones',
    f'Hombres: ${male_sales:,.2f} ({male_sales/total_sales*100:.1f}%) - {len(df[df["Gender"]=="Male"])} transacciones',
    f'Ticket promedio Mujeres: ${df[df["Gender"]=="Female"]["Total"].mean():.2f}',
    f'Ticket promedio Hombres: ${df[df["Gender"]=="Male"]["Total"].mean():.2f}',
    'Gráfica: Gráfico de barras comparativo por género.',
]
for item in slide2_items:
    doc.add_paragraph(item, style='List Bullet')

# Slide 3
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Slide 3: Distribución por Horas del Día')
run.bold = True

slide3_items = [
    'Horario de operación: 10:00 a 21:00 hrs (12 horas).',
    'División equitativa en 3 períodos de 4 horas:',
]
for item in slide3_items:
    doc.add_paragraph(item, style='List Bullet')

morning = df[df['Hour'].between(10, 13)]['Total'].sum()
afternoon = df[df['Hour'].between(14, 17)]['Total'].sum()
evening = df[df['Hour'].between(18, 21)]['Total'].sum()

periods = [
    f'Mañana (10-13h): ${morning:,.2f} ({morning/total_sales*100:.1f}%)',
    f'Tarde (14-17h): ${afternoon:,.2f} ({afternoon/total_sales*100:.1f}%)',
    f'Noche (18-21h): ${evening:,.2f} ({evening/total_sales*100:.1f}%)',
]
for item in periods:
    doc.add_paragraph(item, style='List Bullet 2')

hourly_sales = df.groupby('Hour')['Total'].sum()
peak_hour = hourly_sales.idxmax()
peak_sales = hourly_sales.max()
doc.add_paragraph(f'Hora pico: {peak_hour}:00 hrs con ${peak_sales:,.2f} en ventas.', style='List Bullet')
doc.add_paragraph('Gráfica: Gráfico de líneas mostrando ventas por hora.', style='List Bullet')

# Slide 4
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Slide 4: Ingresos por Línea de Productos')
run.bold = True

product_sales_sorted = df.groupby('Product line')['Total'].sum().sort_values(ascending=False)
for product, sales in product_sales_sorted.items():
    pct = (sales / df['Total'].sum()) * 100
    count = len(df[df['Product line'] == product])
    doc.add_paragraph(f'{product}: ${sales:,.2f} ({pct:.1f}%) - {count} transacciones', style='List Bullet')

doc.add_paragraph('Gráfica: Gráfico de barras horizontales ordenado de mayor a menor.', style='List Bullet')

# Slide 5
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Slide 5: Ingresos por Línea de Productos y Género')
run.bold = True

slide5_items = [
    'Health and beauty: Hombres 62.3% vs Mujeres 37.7% (brecha 24.5%)',
    'Food and beverages: Mujeres 59.1% vs Hombres 40.9% (brecha 18.2%)',
    'Fashion accessories: Mujeres 56.0% vs Hombres 44.0% (brecha 12.1%)',
    'Home and lifestyle: Mujeres 55.8% vs Hombres 44.2% (brecha 11.5%)',
    'Sports and travel: Mujeres 51.8% vs Hombres 48.2% (brecha 3.7%)',
    'Electronic accessories: Hombres 50.1% vs Mujeres 49.9% (brecha 0.2%)',
    'Gráfica: Gráfico de barras apiladas por producto y género.',
]
for item in slide5_items:
    doc.add_paragraph(item, style='List Bullet')

# Slide 6
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Slide 6: Compras por Sucursal')
run.bold = True

for branch in ['A', 'B', 'C']:
    branch_data = df[df['Branch'] == branch]
    sales = branch_data['Total'].sum()
    count = len(branch_data)
    avg = branch_data['Total'].mean()
    city = branch_data['City'].iloc[0]
    pct = (sales / df['Total'].sum()) * 100
    doc.add_paragraph(f'Sucursal {branch} ({city}): ${sales:,.2f} ({pct:.1f}%) - {count} transacciones - Ticket promedio: ${avg:.2f}', style='List Bullet')

doc.add_paragraph('Gráfica: Gráfico de barras por sucursal.', style='List Bullet')

# Slide 7
doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Slide 7: Compras por Línea de Productos')
run.bold = True

transaction_ranking = df.groupby('Product line').size().sort_values(ascending=False)
for i, (product, count) in enumerate(transaction_ranking.items(), 1):
    avg_ticket = df[df['Product line'] == product]['Total'].mean()
    doc.add_paragraph(f'{i}. {product}: {count} transacciones - Ticket promedio: ${avg_ticket:.2f}', style='List Bullet')

doc.add_paragraph('Gráfica: Gráfico de barras horizontales por cantidad de transacciones.', style='List Bullet')

# Slide 8: Principales puntos encontrados
doc.add_heading('3.3 Principales Puntos Encontrados (Slide 8)', level=2)

p = doc.add_paragraph()
p.add_run('3 Oportunidades identificadas para campañas de marketing:').bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Oportunidad 1: Desequilibrio de género en productos')
run.bold = True

opp1_items = [
    'Health and beauty tiene una brecha de 24.5% entre géneros: los hombres representan el 62.3% de las ventas y las mujeres solo el 37.7%.',
    'Food and beverages muestra la situación inversa: mujeres 59.1% vs hombres 40.9% (brecha 18.2%).',
    'El género minoritario en cada producto está subrepresentado, lo que indica un mercado no captado.',
    'Campaña: Atraer al género minoritario a estos productos mediante promociones y comunicación dirigida.',
]
for item in opp1_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Oportunidad 2: Diferencia en comportamiento de clientes Member vs Normal')
run.bold = True

member_avg = df[df['Customer type'] == 'Member']['Total'].mean()
normal_avg = df[df['Customer type'] == 'Normal']['Total'].mean()
diff = member_avg - normal_avg

opp2_items = [
    f'Los clientes Member tienen un ticket promedio de ${member_avg:.2f}, mientras que los Normal de ${normal_avg:.2f}.',
    f'Los Members gastan ${diff:.2f} más por transacción ({(member_avg/normal_avg-1)*100:.1f}% más).',
    'En Health and beauty la diferencia es aún mayor: Members gastan 19.7% más que Normal.',
    f'Hay {len(df[df["Customer type"]=="Normal"])} clientes Normal que podrían convertirse a Member.',
    'Campaña: Programa de fidelización para convertir clientes Normal a Member.',
]
for item in opp2_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('Oportunidad 3: Horarios con baja densidad de ventas')
run.bold = True

avg_hourly = df.groupby('Hour')['Total'].sum().mean()

opp3_items = [
    f'El promedio de ventas por hora es de ${avg_hourly:,.2f}.',
    f'Las 10:00h registran solo ${hourly_sales[10]:,.2f} en ventas (51 transacciones), significativamente por debajo del promedio.',
    f'Las 21:00h registran solo ${hourly_sales[21]:,.2f} en ventas (38 transacciones), también por debajo del promedio.',
    'Estos dos horarios representan las horas de apertura y cierre, con menor afluencia de clientes.',
    'Campaña: Promociones especiales en horarios de apertura y cierre para incentivar la visita en esos momentos.',
]
for item in opp3_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# 4. GRÁFICAS QUE USARON
# ============================================================
doc.add_heading('4. Gráficas Utilizadas', level=1)

doc.add_heading('4.1 ¿Por qué elegimos cada gráfica?', level=2)

graficas = [
    {
        'nombre': 'Gráfico de barras comparativo (Compras por Género)',
        'razon': 'Elegimos un gráfico de barras porque permite comparar de forma directa y clara dos categorías (Mujeres vs Hombres). La diferencia visual entre las barras hace evidente la proporción de ventas de cada género sin necesidad de leer números exactos.'
    },
    {
        'nombre': 'Gráfico de líneas (Distribución por Horas)',
        'razon': 'El gráfico de líneas es ideal para mostrar tendencias a lo largo del tiempo (en este caso, horas del día). Permite ver de un vistazo los picos y valles de actividad, identificando fácilmente la hora pico (15:00h) y los horarios de menor actividad (10:00h y 21:00h).'
    },
    {
        'nombre': 'Gráfico de barras horizontales (Ingresos por Línea de Productos)',
        'razon': 'Las barras horizontales son ideales cuando los nombres de las categorías son largos (como los nombres de líneas de productos). Además, al ordenarlas de mayor a menor, se facilita la comparación y el ranking visual inmediato.'
    },
    {
        'nombre': 'Gráfico de barras apiladas (Ingresos por Producto y Género)',
        'razon': 'Las barras apiladas permiten ver simultáneamente dos dimensiones: el total por producto y la composición por género dentro de cada producto. Esto hace visible la brecha de género en cada línea de productos de forma intuitiva.'
    },
    {
        'nombre': 'Gráfico de barras agrupadas (Compras por Sucursal)',
        'razon': 'Las barras agrupadas permiten comparar las 3 sucursales lado a lado, haciendo evidente las diferencias de rendimiento entre ellas. Es más efectivo que una tabla de números para identificar rápidamente cuál sucursal lidera.'
    },
]

for grafica in graficas:
    p = doc.add_paragraph()
    run = p.add_run(grafica['nombre'])
    run.bold = True
    doc.add_paragraph(grafica['razon'], style='List Bullet')
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 5. CLUTTER
# ============================================================
doc.add_heading('5. Clutter (Limpieza Visual)', level=1)

doc.add_heading('5.1 ¿Qué limpieza hicimos para que las gráficas queden mejor?', level=2)

p = doc.add_paragraph()
p.add_run('Aplicamos los siguientes principios de reducción de clutter para mejorar la claridad de nuestras gráficas:').italic = True

clutter_items = [
    {
        'titulo': 'Eliminación de bordes y marcos innecesarios',
        'detalle': 'Removimos los bordes de las gráficas y las cajas alrededor de las leyendas. Esto reduce el "ruido visual" y permite que los datos sean el foco principal.'
    },
    {
        'titulo': 'Reducción de líneas de grilla',
        'detalle': 'Usamos líneas de grilla suaves (color gris claro, con transparencia) solo en el eje Y cuando es necesario para facilitar la lectura. Eliminamos la grilla del eje X cuando no aporta valor.'
    },
    {
        'titulo': 'Paleta de colores limitada y con propósito',
        'detalle': 'Usamos máximo 2-3 colores por gráfica. Los colores se eligieron con contraste suficiente para diferenciar categorías sin saturar visualmente. Evitamos colores brillantes o neón.'
    },
    {
        'titulo': 'Eliminación de etiquetas redundantes',
        'detalle': 'Si el título de la gráfica ya indica qué se mide, no repetimos esa información en los ejes. Por ejemplo, si el título dice "Ventas por Hora", no agregamos "Ventas ($)" en el eje Y.'
    },
    {
        'titulo': 'Simplificación de números',
        'detalle': 'Redondeamos los valores a números enteros o con un decimal cuando la precisión extra no aporta al análisis. Por ejemplo, "$322.97" en lugar de "$322.96675".'
    },
    {
        'titulo': 'Orden lógico de los datos',
        'detalle': 'Ordenamos las barras de mayor a menor (o en orden cronológico para horas) para facilitar la lectura natural y la comparación entre categorías.'
    },
]

for item in clutter_items:
    p = doc.add_paragraph()
    run = p.add_run(item['titulo'])
    run.bold = True
    doc.add_paragraph(item['detalle'], style='List Bullet')
    doc.add_paragraph('')

doc.add_page_break()

# ============================================================
# 6. FOCO
# ============================================================
doc.add_heading('6. Foco', level=1)

doc.add_heading('6.1 ¿Dónde aplicamos el concepto de Foco en una slide?', level=2)

p = doc.add_paragraph()
run = p.add_run('Slide elegida: Slide 5 - Ingresos por Línea de Productos y Género')
run.bold = True

doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('Aplicación del concepto de Foco:').bold = True

foco_items = [
    {
        'titulo': 'Color como herramienta de foco',
        'detalle': 'En la gráfica de barras apiladas, usamos un color intenso (rojo/naranja) para resaltar Health and beauty, que es el producto con mayor brecha de género (24.5%). El resto de los productos se muestran en tonos grises o apagados. Esto dirige la atención del espectador directamente al hallazgo más importante.'
    },
    {
        'titulo': 'Anotación destacada',
        'detalle': 'Agregamos una anotación con flecha señalando la brecha de 24.5% en Health and beauty, con el texto "Mayor oportunidad de crecimiento". Esta anotación actúa como un llamado visual que guía la interpretación del dato.'
    },
    {
        'titulo': 'Título orientado al insight',
        'detalle': 'En lugar de un título genérico como "Ventas por Producto y Género", usamos un título que comunica el hallazgo: "Health & Beauty: los hombres compran el doble que las mujeres". Esto le dice al espectador qué debe ver antes de mirar la gráfica.'
    },
    {
        'titulo': 'Jerarquía visual',
        'detalle': 'El dato clave (brecha de 24.5%) se muestra en tamaño grande y en negrita junto a la gráfica, mientras que los datos secundarios (otras líneas de productos) se mantienen en tamaño normal. Esto crea una jerarquía clara de información.'
    },
]

for item in foco_items:
    p = doc.add_paragraph()
    run = p.add_run(item['titulo'])
    run.bold = True
    doc.add_paragraph(item['detalle'], style='List Bullet')
    doc.add_paragraph('')

p = doc.add_paragraph()
p.add_run('Resultado: ').bold = True
p.add_run('Al aplicar el concepto de Foco en esta slide, el espectador identifica inmediatamente la oportunidad principal (brecha de género en Health and beauty) sin perderse en los detalles de las otras 5 líneas de productos. El foco transforma una gráfica informativa en una gráfica persuasiva que justifica la primera campaña de marketing.')

# ============================================================
# GUARDAR DOCUMENTO
# ============================================================
output_path = '/Users/javierfernandez/Desktop/Taller/Desafio3_Documento.docx'
doc.save(output_path)
print(f"Documento guardado exitosamente en: {output_path}")
