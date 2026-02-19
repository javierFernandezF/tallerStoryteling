from docx import Document
from docx.shared import Pt, RGBColor
import pandas as pd

# Cargar datos para referencias exactas
def load_data():
    data_rows = []
    with open('/Users/javierfernandez/Desktop/Taller/Libro1.csv', 'r', encoding='utf-8-sig') as f:
        lines = f.readlines()
    header_line = lines[0].strip()
    if header_line.startswith('"') and header_line.endswith('"'):
        header = header_line[1:-1].split(';')
    for line in lines[1:]:
        line = line.strip()
        if line.startswith('"') and line.endswith('"'):
            row_data = line[1:-1].split(';')
            data_rows.append(row_data)
    df = pd.DataFrame(data_rows, columns=header)
    numeric_cols = ['Unit price', 'Quantity', 'Tax 5%', 'Total', 'Time', 'cogs', 'gross margin percentage', 'gross income', 'Rating']
    for col in numeric_cols:
        if col in df.columns:
            df[col] = df[col].str.replace(',', '.').astype(float)
    df['Hour'] = (df['Time'] * 24).round().astype(int).clip(0, 23)
    return df

df = load_data()

# Abrir documento existente
doc = Document('/Users/javierfernandez/Desktop/Taller/Desafio3_Documento.docx')

# Buscar la sección 1.2 y reemplazar el contenido
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.text == "1.2 Principales Observaciones del DataSet":
        # Encontramos el heading, ahora buscamos el contenido después
        start_index = i + 1
        
        # Eliminar párrafos existentes de viñetas
        paragraphs_to_remove = []
        for j in range(start_index, len(doc.paragraphs)):
            if doc.paragraphs[j].style.name == 'List Bullet' or '•' in doc.paragraphs[j].text:
                paragraphs_to_remove.append(doc.paragraphs[j])
            elif doc.paragraphs[j].text.strip() == '' or doc.paragraphs[j].text.startswith('2.'):
                break
        
        # Insertar nuevo contenido narrativo después del heading
        insert_position = start_index
        
        # Párrafo 1: Volumen y alcance
        p1 = doc.paragraphs[insert_position]._element.getparent().insert(insert_position, doc.add_paragraph()._element)
        new_p1 = doc.paragraphs[insert_position]
        new_p1.text = f"El dataset analizado representa un volumen significativo de información comercial, conteniendo exactamente {len(df)} transacciones de ventas registradas durante un período de 3 meses de operación continua de una cadena de supermercados. Este conjunto de datos nos proporciona una visión integral del comportamiento de compra de los clientes y las dinámicas operativas del negocio durante un trimestre completo, lo que constituye una muestra representativa para el análisis de patrones de consumo y la identificación de oportunidades de crecimiento."
        
        # Párrafo 2: Rendimiento financiero
        p2 = doc.paragraphs[insert_position + 1]._element.getparent().insert(insert_position + 1, doc.add_paragraph()._element)
        new_p2 = doc.paragraphs[insert_position + 1]
        new_p2.text = f"Desde una perspectiva financiera, los datos revelan un rendimiento sólido con ventas totales registradas de ${df['Total'].sum():,.2f} durante el período analizado. El ticket promedio por transacción se sitúa en ${df['Total'].mean():.2f}, lo que indica un nivel de gasto consistente por parte de los clientes. Esta cifra promedio sugiere que la cadena ha logrado mantener un equilibrio efectivo entre la accesibilidad de precios y la generación de ingresos por transacción, posicionándose en un rango medio-alto del mercado de supermercados."
        
        # Párrafo 3: Distribución geográfica
        branch_a = len(df[df['Branch'] == 'A'])
        branch_b = len(df[df['Branch'] == 'B'])
        branch_c = len(df[df['Branch'] == 'C'])
        p3 = doc.paragraphs[insert_position + 2]._element.getparent().insert(insert_position + 2, doc.add_paragraph()._element)
        new_p3 = doc.paragraphs[insert_position + 2]
        new_p3.text = f"La distribución geográfica de las operaciones muestra una presencia equilibrada en tres ubicaciones estratégicas. La sucursal A, ubicada en Yangon, registra {branch_a} transacciones, representando la mayor actividad comercial. La sucursal B en Mandalay presenta {branch_b} transacciones, mientras que la sucursal C en Naypyitaw registra {branch_c} transacciones. Esta distribución relativamente uniforme entre las tres ubicaciones sugiere una estrategia de expansión geográfica exitosa y una capacidad operativa consistente across diferentes mercados urbanos."
        
        # Párrafo 4: Perfil demográfico
        female_count = len(df[df['Gender'] == 'Female'])
        male_count = len(df[df['Gender'] == 'Male'])
        member_count = len(df[df['Customer type'] == 'Member'])
        normal_count = len(df[df['Customer type'] == 'Normal'])
        p4 = doc.paragraphs[insert_position + 3]._element.getparent().insert(insert_position + 3, doc.add_paragraph()._element)
        new_p4 = doc.paragraphs[insert_position + 3]
        new_p4.text = f"El perfil demográfico de la clientela revela características interesantes en términos de composición y fidelización. La distribución por género es notablemente equilibrada, con {female_count} transacciones realizadas por mujeres y {male_count} por hombres, lo que indica que la cadena ha logrado atraer de manera efectiva a ambos segmentos demográficos. En cuanto a la tipología de clientes, existe una división casi perfecta entre clientes Member ({member_count} transacciones) y clientes Normal ({normal_count} transacciones), sugiriendo que el programa de membresía ha alcanzado una penetración del 50% en la base de clientes activos."
        
        # Párrafo 5: Diversificación de productos y operaciones
        product_lines = df['Product line'].nunique()
        p5 = doc.paragraphs[insert_position + 4]._element.getparent().insert(insert_position + 4, doc.add_paragraph()._element)
        new_p5 = doc.paragraphs[insert_position + 4]
        new_p5.text = f"La diversificación del portafolio de productos abarca {product_lines} líneas principales: Health and beauty, Electronic accessories, Home and lifestyle, Sports and travel, Food and beverages, y Fashion accessories. Esta variedad demuestra una estrategia comercial integral que va más allá del concepto tradicional de supermercado, incorporando categorías de productos que atienden diferentes necesidades y momentos de consumo de los clientes. Las operaciones se desarrollan en un horario extendido de 12 horas diarias, desde las 10:00 hasta las 21:00 hrs, maximizando las oportunidades de venta y adaptándose a los diferentes patrones de compra de los consumidores. Adicionalmente, la cadena ofrece flexibilidad en los métodos de pago, aceptando tres modalidades principales: Ewallet, Cash y Credit card, lo que facilita la experiencia de compra y se adapta a las preferencias de pago de diferentes segmentos de clientes."
        
        break

# Guardar documento modificado
doc.save('/Users/javierfernandez/Desktop/Taller/Desafio3_Documento_Mejorado.docx')
print("Documento mejorado guardado como: Desafio3_Documento_Mejorado.docx")
