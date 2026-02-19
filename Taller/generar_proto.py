from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

# Cargar datos
data_rows = []
with open('/Users/javierfernandez/Desktop/Taller/Libro1.csv', 'r', encoding='utf-8-sig') as f:
    lines = f.readlines()
header_line = lines[0].strip()
if header_line.startswith('"') and header_line.endswith('"'):
    header = header_line[1:-1].split(';')
for line in lines[1:]:
    line = line.strip()
    if line.startswith('"') and line.endswith('"'):
        row_data = line[1:-1].split(';')
        data_rows.append(row_data)
df = pd.DataFrame(data_rows, columns=header)
numeric_cols = ['Unit price', 'Quantity', 'Tax 5%', 'Total', 'Time', 'cogs', 'gross margin percentage', 'gross income', 'Rating']
for col in numeric_cols:
    if col in df.columns:
        df[col] = df[col].str.replace(',', '.').astype(float)
df['Hour'] = (df['Time'] * 24).round().astype(int).clip(0, 23)

# Crear documento
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ---- PORTADA simple ----
for _ in range(5):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Desafío 3 - Supermarket Sales')
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph('')
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Análisis de Datos - Borrador')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_paragraph('')
fecha = doc.add_paragraph()
fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fecha.add_run('11 de febrero de 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ============================================================
# 1. DESCRIPCIÓN DEL DATASET - Completa
# ============================================================
doc.add_heading('1. Descripción del DataSet', level=1)

doc.add_heading('Tamaño y Campos', level=2)

p = doc.add_paragraph()
p.add_run('Tamaño: ').bold = True
p.add_run(f'{len(df)} registros, {len(df.columns)} campos.')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Campos:').bold = True

# Tabla de campos
fields = [
    ('Invoice ID', 'Texto', 'ID de factura'),
    ('Branch', 'Texto', 'Sucursal (A, B, C)'),
    ('City', 'Texto', 'Ciudad'),
    ('Customer type', 'Texto', 'Member o Normal'),
    ('Gender', 'Texto', 'Género'),
    ('Product line', 'Texto', 'Línea de producto'),
    ('Unit price', 'Numérico', 'Precio unitario'),
    ('Quantity', 'Numérico', 'Cantidad'),
    ('Tax 5%', 'Numérico', 'Impuesto'),
    ('Total', 'Numérico', 'Total de la transacción'),
    ('Date', 'Fecha', 'Fecha'),
    ('Time', 'Numérico', 'Hora'),
    ('Payment', 'Texto', 'Método de pago'),
    ('cogs', 'Numérico', 'Costo de bienes'),
    ('gross margin percentage', 'Numérico', 'Margen bruto %'),
    ('gross income', 'Numérico', 'Ingreso bruto'),
    ('Rating', 'Numérico', 'Calificación (1-10)'),
]

table = doc.add_table(rows=len(fields)+1, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(['Campo', 'Tipo', 'Descripción']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for par in cell.paragraphs:
        for run in par.runs:
            run.bold = True

for i, (f, t, d) in enumerate(fields):
    table.rows[i+1].cells[0].text = f
    table.rows[i+1].cells[1].text = t
    table.rows[i+1].cells[2].text = d

doc.add_paragraph('')

doc.add_heading('Principales Observaciones', level=2)

obs = [
    f'1000 transacciones en 3 sucursales.',
    f'Ventas totales: ${df["Total"].sum():,.2f}.',
    f'Ticket promedio: ${df["Total"].mean():.2f}.',
    f'Distribución de género equilibrada: {len(df[df["Gender"]=="Female"])} mujeres, {len(df[df["Gender"]=="Male"])} hombres.',
    f'6 líneas de productos.',
    f'Horario de operación: 10:00 a 21:00 hrs.',
]
for o in obs:
    doc.add_paragraph(o, style='List Bullet')

doc.add_page_break()

# ============================================================
# 2. CONTEXTO - Completa
# ============================================================
doc.add_heading('2. Contexto', level=1)

doc.add_heading('¿Quién?', level=2)
p = doc.add_paragraph()
p.add_run('Cliente: ').bold = True
p.add_run('Cadena de supermercados con 3 sucursales:')

sucursales = [
    f'Sucursal A - Yangon ({len(df[df["Branch"]=="A"])} transacciones)',
    f'Sucursal B - Mandalay ({len(df[df["Branch"]=="B"])} transacciones)',
    f'Sucursal C - Naypyitaw ({len(df[df["Branch"]=="C"])} transacciones)',
]
for s in sucursales:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('¿Qué?', level=2)
p = doc.add_paragraph('La empresa necesita realizar 3 campañas de Marketing (una por mes) para maximizar los ingresos por ventas. Para eso necesitamos analizar los datos y encontrar oportunidades.')

doc.add_page_break()

# ============================================================
# 3. STORYBOARD - Parcialmente completo
# ============================================================
doc.add_heading('3. Storyboard', level=1)

# Nota de "en progreso"
p = doc.add_paragraph()
run = p.add_run('[EN PROGRESO - Definiendo estructura de slides]')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

doc.add_paragraph('')

# Slide 1 - hecho
doc.add_heading('Slide 1: Introducción', level=2)
intro_items = [
    'Título: "Análisis de Ventas - Cadena de Supermercados"',
    f'Datos: {len(df)} transacciones, ${df["Total"].sum():,.2f} en ventas, 3 meses.',
    'Pregunta central: ¿Dónde están las oportunidades para aumentar las ventas?',
]
for item in intro_items:
    doc.add_paragraph(item, style='List Bullet')

# Slide 2 - hecho
doc.add_heading('Slide 2: Compras por Género', level=2)
female_sales = df[df['Gender'] == 'Female']['Total'].sum()
male_sales = df[df['Gender'] == 'Male']['Total'].sum()
total_sales = female_sales + male_sales

items = [
    f'Mujeres: ${female_sales:,.2f} ({female_sales/total_sales*100:.1f}%) - {len(df[df["Gender"]=="Female"])} transacciones',
    f'Hombres: ${male_sales:,.2f} ({male_sales/total_sales*100:.1f}%) - {len(df[df["Gender"]=="Male"])} transacciones',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Slide 3 - hecho
doc.add_heading('Slide 3: Distribución por Horas', level=2)
morning = df[df['Hour'].between(10, 13)]['Total'].sum()
afternoon = df[df['Hour'].between(14, 17)]['Total'].sum()
evening = df[df['Hour'].between(18, 21)]['Total'].sum()

items = [
    'Horario: 10:00 a 21:00 hrs',
    f'Mañana (10-13h): ${morning:,.2f} ({morning/total_sales*100:.1f}%)',
    f'Tarde (14-17h): ${afternoon:,.2f} ({afternoon/total_sales*100:.1f}%)',
    f'Noche (18-21h): ${evening:,.2f} ({evening/total_sales*100:.1f}%)',
    f'Hora pico: 15:00 hrs',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# Slides 4-7 - pendientes con notas
doc.add_heading('Slide 4: Ingresos por Línea de Productos', level=2)
p = doc.add_paragraph()
run = p.add_run('[TODO: Armar con datos del análisis]')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_heading('Slide 5: Ingresos por Producto y Género', level=2)
p = doc.add_paragraph()
run = p.add_run('[TODO: Cruzar producto con género]')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_heading('Slide 6: Compras por Sucursal', level=2)
p = doc.add_paragraph()
run = p.add_run('[TODO: Comparar sucursales A, B, C]')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_heading('Slide 7: Compras por Línea de Productos', level=2)
p = doc.add_paragraph()
run = p.add_run('[TODO: Ranking de productos por transacciones]')
run.font.color.rgb = RGBColor(255, 0, 0)

# Slide 8 - pendiente
doc.add_heading('Slide 8: Principales Puntos / Oportunidades', level=2)
p = doc.add_paragraph()
run = p.add_run('[TODO: Definir las 3 oportunidades para campañas]')
run.font.color.rgb = RGBColor(255, 0, 0)

doc.add_page_break()

# ============================================================
# 4. GRÁFICAS - Solo títulos y notas
# ============================================================
doc.add_heading('4. Gráficas', level=1)

p = doc.add_paragraph()
run = p.add_run('[PENDIENTE - Estamos definiendo qué gráficas usar]')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Ideas iniciales:').bold = True

graficas_ideas = [
    'Barras comparativas para género',
    'Líneas para distribución horaria',
    'Barras horizontales para productos (nombres largos)',
]
for g in graficas_ideas:
    doc.add_paragraph(g, style='List Bullet')

doc.add_paragraph('')

# ============================================================
# 5. CLUTTER - Vacío con nota
# ============================================================
doc.add_heading('5. Clutter', level=1)
p = doc.add_paragraph()
run = p.add_run('[PENDIENTE - Completar cuando tengamos las gráficas armadas]')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

doc.add_paragraph('')

# ============================================================
# 6. FOCO - Vacío con nota
# ============================================================
doc.add_heading('6. Foco', level=1)
p = doc.add_paragraph()
run = p.add_run('[PENDIENTE - Elegir slide donde aplicar foco]')
run.font.color.rgb = RGBColor(255, 0, 0)
run.bold = True

# Guardar
output_path = '/Users/javierfernandez/Desktop/Taller/Desafio3_Borrador.docx'
doc.save(output_path)
print(f"Proto-versión guardada en: {output_path}")
